import os
import io
import datetime
import json
import base64
import pandas as pd
import streamlit as st
from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage, ToolMessage
from langchain_core.tools import tool
from langchain_core.documents import Document as LangchainDocument
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from ddgs import DDGS

# ==========================================
# 0. 核心配置区
# ==========================================
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

# ==========================================
# 1. 页面配置
# ==========================================
st.set_page_config(page_title="全能法规智能助理", page_icon="🤖", layout="wide")


# ==========================================
# 2. 轻量级身份验证系统
# ==========================================
def check_password():
    USER_CREDENTIALS = {"admin": "123456", "boss": "888888"}

    def password_entered():
        username = st.session_state["login_username"]
        password = st.session_state["login_password"]
        api_key_input = st.session_state["login_api_key"].strip()

        if username in USER_CREDENTIALS and USER_CREDENTIALS[username] == password:
            if not api_key_input:
                st.session_state["password_correct"] = False
                st.session_state["login_error"] = "🚫 请输入有效的 DeepSeek API Key！"
            else:
                st.session_state["password_correct"] = True
                st.session_state["deepseek_api_key"] = api_key_input
                del st.session_state["login_password"]
                if "login_error" in st.session_state:
                    del st.session_state["login_error"]
        else:
            st.session_state["password_correct"] = False
            st.session_state["login_error"] = "🚫 用户名或密码错误，请重试！"

    if "password_correct" not in st.session_state or not st.session_state["password_correct"]:
        st.title("🔒 欢迎访问全能智能 Agent")
        st.info("请输入您的账号、密码以及 DeepSeek API Key 以继续。")
        st.text_input("用户名", key="login_username")
        st.text_input("密码", type="password", key="login_password")
        st.text_input("DeepSeek API Key (sk-...)", type="password", key="login_api_key")
        st.button("登录", on_click=password_entered, type="primary")
        if "login_error" in st.session_state:
            st.error(st.session_state["login_error"])
        return False
    return True


# ==========================================
# 🌟 核心拦截器
# ==========================================
if check_password():

    # ==========================================
    # 3. 性能优化与工具定义
    # ==========================================
    @st.cache_resource
    def load_embedding_model():
        return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")


    @tool
    def modify_word_document(action: str, original_text: str, revised_text: str, comment: str) -> str:
        """【文档修改工具】当用户同意修改方案时调用此工具。
        参数说明：
        - action (str): 必须选 "replace" 或 "append"。"replace" 用于修改或替换文档中已有的段落；"append" 用于在文档【最末尾】追加完全缺失的新章节。
        - original_text (str): 需要被替换的原文片段。如果 action 选了 "append"，此项请填空字符串 ""。
        - revised_text (str): 新的文本内容。
        - comment (str): 给用户的批注或说明。
        """
        if "current_file_path" not in st.session_state or not st.session_state.current_file_path:
            return "操作失败：当前没有加载任何文档。"
        source_path = st.session_state.current_file_path
        try:
            doc = DocxDocument(source_path)
            modified_count = 0

            # 🌟 新增逻辑：如果是追加章节，直接写在文档末尾
            if action == "append":
                doc.add_paragraph("")  # 加个空行隔开
                new_p = doc.add_paragraph()
                run_text = new_p.add_run(revised_text)
                run_text.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run_comment = new_p.add_run(f" [AI 追加章节: {comment}]")
                run_comment.font.color.rgb = RGBColor(255, 0, 0)
                modified_count += 1

            # 原有的逻辑：查找并替换现有文本
            else:
                for paragraph in doc.paragraphs:
                    if original_text and original_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(original_text, revised_text)
                        for run in paragraph.runs:
                            if revised_text in run.text:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        run_comment = paragraph.add_run(f" [AI 批注: {comment}]")
                        run_comment.font.color.rgb = RGBColor(255, 0, 0)
                        modified_count += 1

            if modified_count > 0:
                doc.save(source_path)  # 叠甲保存

                # 写入纯内存数据流并生成下载按钮
                word_buffer = io.BytesIO()
                doc.save(word_buffer)
                word_bytes = word_buffer.getvalue()
                b64_data = base64.b64encode(word_bytes).decode()

                st.session_state.latest_word_b64 = b64_data
                st.session_state.newly_modified_word_trigger = True
                return f"成功：文档已{'在末尾追加新章节' if action == 'append' else '修改完毕'}。请回复用户：'文档已更新，请点击下方的蓝色按钮直接下载修订版。'"

            return f"失败：未在原文档中找到原文片段 '{original_text}'。"
        except Exception as e:
            return f"修改文档时发生系统错误: {str(e)}"


    @tool
    def search_document_content(query: str) -> str:
        """【本地文档检索工具】当用户问起“文档里是怎么写的”时调用。"""
        if "vector_db" not in st.session_state or st.session_state.vector_db is None:
            return "操作失败：文档尚未被解析为向量数据库。"
        try:
            docs = st.session_state.vector_db.similarity_search(query, k=3)
            if not docs: return "检索完毕：没有找到相关内容。"
            context = "\n\n".join([f"原文段落 {i + 1}: {d.page_content}" for i, d in enumerate(docs)])
            return f"【本地检索结果】\n{context}"
        except Exception as e:
            return f"检索文档时发生错误: {str(e)}"


    @tool
    def search_latest_medical_regulations(query: str, time_limit: str = "m") -> str:
        """【互联网实时搜索工具】获取当前最新的法规、新闻。"""
        try:
            results = DDGS().text(query, max_results=3, timelimit=time_limit)
            if not results:
                return f"在 {time_limit} 范围内未能找到关于 '{query}' 的最新信息。"
            formatted_results = "\n\n".join(
                [f"标题: {res['title']}\n摘要: {str(res.get('body', ''))[:150]}...\n链接: {res['href']}" for res in
                 results])
            return f"【实时联网搜索结果】\n{formatted_results}"
        except Exception as e:
            return f"联网搜索发生异常: {str(e)}"


    @tool
    def generate_excel_matrix(json_data: str) -> str:
        """【Excel生成工具】生成NC整改矩阵。"""
        try:
            data = json.loads(json_data)
            df = pd.DataFrame(data)

            excel_buffer = io.BytesIO()
            df.to_excel(excel_buffer, index=False)
            excel_bytes = excel_buffer.getvalue()
            b64_data = base64.b64encode(excel_bytes).decode()

            st.session_state.latest_excel_b64 = b64_data
            st.session_state.newly_generated_excel_trigger = True

            return "成功：Excel矩阵已在内存中生成。请回复用户：'表格已为您生成，请点击下方的绿色按钮直接下载。'"
        except Exception as e:
            return f"生成 Excel 时发生错误: {str(e)}。请检查传入的 JSON 格式。"


    @tool
    def get_file_download_link(file_type: str) -> str:
        """【获取文件下载链接】当用户主动要求“下载Word文档”、“给我Excel链接”等需求时，调用此工具。"""
        if file_type.lower() == 'word':
            if st.session_state.get("latest_word_b64"):
                st.session_state.newly_modified_word_trigger = True
                return "已触发 Word 下载按钮，请告诉用户：'文档下载链接已为您重新生成，请点击下方蓝色按钮获取。'"
            return "当前没有内存中的 Word 文档可供下载。"

        elif file_type.lower() == 'excel':
            if st.session_state.get("latest_excel_b64"):
                st.session_state.newly_generated_excel_trigger = True
                return "已触发 Excel 下载按钮，请告诉用户：'表格下载链接已为您重新生成，请点击下方绿色按钮获取。'"
            return "当前没有内存中的 Excel 表格可供下载。"

        return "未知的文件类型。请仅请求 word 或 excel。"


    @tool
    def update_task_board(content: str) -> str:
        """【全局备忘录工具】当你分析长文档得出多个修改意见、或需要记住待办事项清单时，调用此工具将其记录在侧边栏，防止遗忘。"""
        st.session_state.task_board = content
        return "成功：已将内容安全写入全局备忘录。请停止修改并询问用户先执行备忘录里的哪一条。"


    AVAILABLE_TOOLS = {
        "modify_word_document": modify_word_document,
        "search_document_content": search_document_content,
        "search_latest_medical_regulations": search_latest_medical_regulations,
        "generate_excel_matrix": generate_excel_matrix,
        "get_file_download_link": get_file_download_link,
        "update_task_board": update_task_board
    }


    def process_document_to_vector_db(file_path):
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if len(p.text.strip()) > 5]
        if not paragraphs: return None
        docs = [LangchainDocument(page_content=text) for text in paragraphs]
        embeddings = load_embedding_model()
        vector_db = FAISS.from_documents(docs, embeddings)
        return vector_db


    # ==========================================
    # 4. 状态初始化与主界面布局
    # ==========================================
    st.title("🤖 医疗器械法规 Agent")

    if "current_file_path" not in st.session_state: st.session_state.current_file_path = None
    if "current_file_name" not in st.session_state: st.session_state.current_file_name = None
    if "vector_db" not in st.session_state: st.session_state.vector_db = None
    if "task_board" not in st.session_state: st.session_state.task_board = ""

    if "latest_word_b64" not in st.session_state: st.session_state.latest_word_b64 = None
    if "latest_excel_b64" not in st.session_state: st.session_state.latest_excel_b64 = None
    if "newly_generated_excel_trigger" not in st.session_state: st.session_state.newly_generated_excel_trigger = False
    if "newly_modified_word_trigger" not in st.session_state: st.session_state.newly_modified_word_trigger = False

    # ==========================================
    # 5. 侧边栏
    # ==========================================
    with st.sidebar:
        st.header("⚙️ 系统状态")
        st.success("✅ 核心 AI 引擎已连接")
        st.markdown("---")
        st.header("📂 文档管理")
        uploaded_file = st.file_uploader("上传待修改的技术文档 (.docx)", type=["docx"])

        if uploaded_file is not None:
            if st.session_state.current_file_name != uploaded_file.name:
                save_path = f"temp_{uploaded_file.name}"
                with open(save_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                st.session_state.current_file_path = save_path
                st.session_state.current_file_name = uploaded_file.name
                with st.spinner("正在解析文档并建立大脑索引..."):
                    st.session_state.vector_db = process_document_to_vector_db(save_path)
                st.success(f"已加载并解析: {uploaded_file.name}")
        else:
            st.session_state.current_file_path = None
            st.session_state.current_file_name = None
            st.session_state.vector_db = None
            st.session_state.latest_word_b64 = None
            st.session_state.latest_excel_b64 = None

        if st.session_state.latest_word_b64:
            st.markdown("---")
            word_bytes = base64.b64decode(st.session_state.latest_word_b64)
            filename = f"revised_{st.session_state.current_file_name}"
            st.download_button("📥 左侧备用下载：修订版 Word", data=word_bytes, file_name=filename,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               type="primary")

        if st.session_state.task_board:
            st.markdown("---")
            st.header("📋 AI 待办备忘录")
            st.info(st.session_state.task_board)
            if st.button("🗑️ 清空备忘录"):
                st.session_state.task_board = ""
                st.rerun()

        st.markdown("---")
        if st.button("🚪 退出登录"):
            del st.session_state["password_correct"]
            if "deepseek_api_key" in st.session_state:
                del st.session_state["deepseek_api_key"]
            st.rerun()

    if st.session_state.current_file_name:
        st.info(f"📄 **当前文档:** `{st.session_state.current_file_name}` | 🧠 **阅读状态:** 已载入记忆。")
    else:
        st.info("💡 **当前未加载文档。** 你可以把我当做全能助手进行日常闲聊，或上传文档进行深度分析。")

    # ==========================================
    # 6. Agent 聊天与核心调度逻辑
    # ==========================================
    real_today = datetime.datetime.now().strftime("%Y年%m月%d日")
    doc_status = f"已在后台系统中加载并解析了文档【{st.session_state.current_file_name}】" if st.session_state.current_file_name else "当前未加载任何文档"
    current_board = st.session_state.task_board if st.session_state.task_board else "暂无内容"

    system_prompt = f"""你是一个全能的AI智能助理，核心专长是资深医疗器械合规专家。
    【重要时间认知】：今天是真实的 {real_today}。
    【当前文档状态】：{doc_status}
    【全局备忘录内容】：{current_board}

    你的工作模式：
    1. 【常规闲聊】：解答通用知识。
    2. 【联网搜索】：调用 `search_latest_medical_regulations` 获取最新信息。
    3. 【文档处理】：调用本地文档搜索及 Word 修改工具。
       ⚠️ 高级计划指令（防死循环）：当你阅读文档并发现【多个错误】需要修改时，你必须先调用 `update_task_board` 工具，把发现的错误清单写进备忘录里！然后停止执行，询问用户：“我已经将发现的问题记录在了左侧的备忘录中，请问我们先从哪一条开始修改？”。绝不要试图一次性修改所有内容以免超时！
       ⚠️ 强制指令：当用户同意你的修改方案时，你必须直接调用 `modify_word_document` 工具执行写入。
    4. 【生成表格】：调用 `generate_excel_matrix` 生成 Excel 文件。
    5. 【召唤链接】：如果用户要求重新获取已经生成的文档或表格下载链接，调用 `get_file_download_link` 工具。
    """

    if "messages" not in st.session_state:
        st.session_state.messages = [
            SystemMessage(content=system_prompt),
            AIMessage(
                content="你好！我已修复了多次修改被覆盖的问题。现在你可以持续对文档下达修改指令，所有的修改都会完整叠加记录并随时供您下载！")
        ]
    else:
        if len(st.session_state.messages) > 0 and isinstance(st.session_state.messages[0], SystemMessage):
            st.session_state.messages[0] = SystemMessage(content=system_prompt)

    user_api_key = st.session_state.get("deepseek_api_key", "")
    llm = ChatOpenAI(api_key=user_api_key, base_url="https://api.deepseek.com", model="deepseek-chat", temperature=0.3,
                     streaming=True, max_retries=3, timeout=60.0)
    llm_with_tools = llm.bind_tools(list(AVAILABLE_TOOLS.values()))

    for msg in st.session_state.messages:
        if isinstance(msg, SystemMessage) or isinstance(msg, ToolMessage): continue
        if isinstance(msg, AIMessage) and not msg.content: continue
        role = "user" if isinstance(msg, HumanMessage) else "assistant"
        with st.chat_message(role):
            st.markdown(msg.content, unsafe_allow_html=True)

    if user_input := st.chat_input("输入指令..."):

        if len(st.session_state.messages) > 41:
            tail = st.session_state.messages[-40:]
            while tail and not isinstance(tail[0], HumanMessage):
                tail.pop(0)
            st.session_state.messages = [st.session_state.messages[0]] + tail

        if len(st.session_state.messages) > 0:
            last_msg = st.session_state.messages[-1]
            if isinstance(last_msg, AIMessage) and getattr(last_msg, "tool_calls", None):
                st.session_state.messages[-1] = AIMessage(content=last_msg.content)

        with st.chat_message("user"):
            st.markdown(user_input, unsafe_allow_html=True)
        st.session_state.messages.append(HumanMessage(content=user_input))

        with st.chat_message("assistant"):
            max_loops = 20
            current_loop = 0

            while current_loop < max_loops:
                current_loop += 1

                stop_btn_container = st.empty()
                message_placeholder = st.empty()
                ai_msg_chunk = None

                status_text = "AI 正在思考..." if current_loop == 1 else "AI 正在分析执行结果..."

                with stop_btn_container.container():
                    if st.button("⏹️ 停止生成", key=f"stop_{len(st.session_state.messages)}_{current_loop}"):
                        st.stop()

                with st.spinner(status_text):
                    current_messages = st.session_state.messages.copy()
                    st.session_state.messages.append(AIMessage(content=""))
                    current_msg_idx = len(st.session_state.messages) - 1

                    for chunk in llm_with_tools.stream(current_messages):
                        if ai_msg_chunk is None:
                            ai_msg_chunk = chunk
                        else:
                            ai_msg_chunk = ai_msg_chunk + chunk

                        st.session_state.messages[current_msg_idx] = ai_msg_chunk

                        if ai_msg_chunk.content:
                            message_placeholder.markdown(ai_msg_chunk.content + " ▌", unsafe_allow_html=True)

                if ai_msg_chunk and ai_msg_chunk.content:
                    message_placeholder.markdown(ai_msg_chunk.content, unsafe_allow_html=True)

                stop_btn_container.empty()

                if not ai_msg_chunk or not ai_msg_chunk.tool_calls:
                    break

                for tool_call in ai_msg_chunk.tool_calls:
                    tool_name = tool_call["name"]
                    tool_args = tool_call["args"]
                    tool_func = AVAILABLE_TOOLS.get(tool_name)

                    try:
                        if tool_func:
                            result_msg = tool_func.invoke(tool_args)
                        else:
                            result_msg = f"系统错误：找不到名为 {tool_name} 的工具。"
                    except Exception as e:
                        result_msg = f"工具执行异常: {str(e)}"

                    tool_message = ToolMessage(content=str(result_msg), tool_call_id=tool_call["id"])
                    st.session_state.messages.append(tool_message)

            if current_loop >= max_loops and ai_msg_chunk and ai_msg_chunk.tool_calls:
                st.warning("⚠️ 思考及搜索轮数达到上限，已强制中断推导。")

            # ==========================================
            # 🌟 拦截器：渲染纯内存下载按钮
            # ==========================================

            if st.session_state.get("newly_generated_excel_trigger") and st.session_state.get("latest_excel_b64"):
                b64 = st.session_state.latest_excel_b64
                filename = "NC_Rectification_Matrix.xlsx"

                html_link = f'<div style="margin-top: 15px;"><a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{filename}" target="_blank" style="display: inline-block; padding: 10px 20px; background-color: #00A67E; color: white; text-align: center; text-decoration: none; border-radius: 6px; font-weight: bold;">📊 点击这里直接下载 Excel 矩阵表</a></div>'
                st.markdown(html_link, unsafe_allow_html=True)

                if len(st.session_state.messages) > 0 and isinstance(st.session_state.messages[-1], AIMessage):
                    st.session_state.messages[-1].content += "\n\n" + html_link

                st.session_state.newly_generated_excel_trigger = False

            if st.session_state.get("newly_modified_word_trigger") and st.session_state.get("latest_word_b64"):
                b64 = st.session_state.latest_word_b64
                filename = f"revised_{st.session_state.current_file_name}" if st.session_state.current_file_name else "revised_document.docx"

                html_link = f'<div style="margin-top: 15px;"><a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" target="_blank" style="display: inline-block; padding: 10px 20px; background-color: #0052CC; color: white; text-align: center; text-decoration: none; border-radius: 6px; font-weight: bold;">📝 点击这里直接下载修订版 Word</a></div>'
                st.markdown(html_link, unsafe_allow_html=True)

                if len(st.session_state.messages) > 0 and isinstance(st.session_state.messages[-1], AIMessage):
                    st.session_state.messages[-1].content += "\n\n" + html_link

                st.session_state.newly_modified_word_trigger = False
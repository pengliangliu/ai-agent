import os
# 【新增】：强制 HuggingFace 使用国内镜像源下载模型
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
import streamlit as st
from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage, ToolMessage
from langchain_core.tools import tool
from langchain_core.documents import Document as LangchainDocument
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

# ==========================================
# 0. 页面配置
# ==========================================
st.set_page_config(page_title="法规文档智能助理", page_icon="🤖", layout="wide")


# ==========================================
# 1. 轻量级身份验证系统
# ==========================================
def check_password():
    USER_CREDENTIALS = {
        "admin": "123456",
        "boss": "888888"
    }

    def password_entered():
        username = st.session_state["login_username"]
        password = st.session_state["login_password"]

        if username in USER_CREDENTIALS and USER_CREDENTIALS[username] == password:
            st.session_state["password_correct"] = True
            del st.session_state["login_password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.title("🔒 欢迎访问医疗法规 Agent")
        st.info("请输入合法的用户名和密码以继续。")
        st.text_input("用户名", key="login_username")
        st.text_input("密码", type="password", key="login_password")
        st.button("登录", on_click=password_entered, type="primary")
        return False

    elif not st.session_state["password_correct"]:
        st.title("🔒 欢迎访问医疗法规 Agent")
        st.text_input("用户名", key="login_username")
        st.text_input("密码", type="password", key="login_password")
        st.button("登录", on_click=password_entered, type="primary")
        st.error("🚫 用户名或密码错误，请重试！")
        return False

    return True


# ==========================================
# 🌟 核心拦截器
# ==========================================
if check_password():

    # ==========================================
    # 2. 性能优化与工具定义
    # ==========================================
    @st.cache_resource
    def load_embedding_model():
        return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")


    @tool
    def modify_word_document(original_text: str, revised_text: str, comment: str) -> str:
        """
        【文档修改工具】当用户同意修改方案，或明确要求“去修改Word文档”、“执行写入”时调用此工具。
        """
        if "current_file_path" not in st.session_state or not st.session_state.current_file_path:
            return "操作失败：当前没有加载任何文档。请告诉用户需要先上传文档才能执行修改。"

        source_path = st.session_state.current_file_path
        output_path = f"revised_{st.session_state.current_file_name}"

        try:
            doc = DocxDocument(source_path)
            modified_count = 0
            for paragraph in doc.paragraphs:
                if original_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(original_text, revised_text)
                    for run in paragraph.runs:
                        if revised_text in run.text:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run_comment = paragraph.add_run(f" [AI 批注: {comment}]")
                    run_comment.font.color.rgb = RGBColor(255, 0, 0)
                    modified_count += 1

            doc.save(output_path)
            if modified_count > 0:
                st.session_state.output_file_path = output_path
                return f"成功：已将修改写入文档并另存为 {output_path}。"
            else:
                return f"失败：未在原文档中找到原文片段 '{original_text}'。"
        except Exception as e:
            return f"修改文档时发生系统错误: {str(e)}"


    @tool
    def search_document_content(query: str) -> str:
        """
        【文档阅读与检索工具】当用户问起“文档里是怎么写的”、“请帮我看看原文档的内容”或者你需要结合上下文制定修改方案时，必须调用此工具。
        """
        if "vector_db" not in st.session_state or st.session_state.vector_db is None:
            return "操作失败：文档尚未被解析为向量数据库，无法检索。请告诉用户需先上传文档。"

        try:
            docs = st.session_state.vector_db.similarity_search(query, k=3)
            if not docs:
                return "检索完毕：在文档中没有找到与该查询相关的内容。"

            context = "\n\n".join([f"原文段落 {i + 1}: {d.page_content}" for i, d in enumerate(docs)])
            return f"【检索结果】\n{context}"
        except Exception as e:
            return f"检索文档时发生错误: {str(e)}"


    AVAILABLE_TOOLS = {
        "modify_word_document": modify_word_document,
        "search_document_content": search_document_content
    }


    def process_document_to_vector_db(file_path):
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if len(p.text.strip()) > 5]
        if not paragraphs:
            return None

        docs = [LangchainDocument(page_content=text) for text in paragraphs]
        embeddings = load_embedding_model()
        vector_db = FAISS.from_documents(docs, embeddings)
        return vector_db


    # ==========================================
    # 3. 状态初始化与主界面布局
    # ==========================================
    st.title("🤖 医疗器械法规 Agent (全能版)")

    if "current_file_path" not in st.session_state:
        st.session_state.current_file_path = None
    if "current_file_name" not in st.session_state:
        st.session_state.current_file_name = None
    if "output_file_path" not in st.session_state:
        st.session_state.output_file_path = None
    if "vector_db" not in st.session_state:
        st.session_state.vector_db = None

    # ==========================================
    # 4. 侧边栏
    # ==========================================
    with st.sidebar:
        st.header("⚙️ 系统设置")
        api_key = st.text_input("输入 DeepSeek API Key", type="password")

        st.markdown("---")
        st.header("📂 文档管理")

        uploaded_file = st.file_uploader("上传待修改的技术文档 (.docx)", type=["docx"])

        if uploaded_file is not None:
            if st.session_state.current_file_name != uploaded_file.name:
                save_path = f"temp_{uploaded_file.name}"
                with open(save_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())

                st.session_state.current_file_path = save_path
                st.session_state.current_file_name = uploaded_file.name

                with st.spinner("正在解析文档并建立大脑索引..."):
                    st.session_state.vector_db = process_document_to_vector_db(save_path)

                st.success(f"已加载并解析: {uploaded_file.name}")
        else:
            st.session_state.current_file_path = None
            st.session_state.current_file_name = None
            st.session_state.output_file_path = None
            st.session_state.vector_db = None

        if st.session_state.output_file_path and os.path.exists(st.session_state.output_file_path):
            st.markdown("---")
            with open(st.session_state.output_file_path, "rb") as file:
                st.download_button("📥 下载修订版 Word", data=file, file_name=st.session_state.output_file_path,
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   type="primary")

        st.markdown("---")
        if st.button("🚪 退出登录"):
            del st.session_state["password_correct"]
            st.rerun()

    # ==========================================
    # 5. 主界面提示
    # ==========================================
    if st.session_state.current_file_name:
        st.info(f"📄 **当前文档:** `{st.session_state.current_file_name}` | 🧠 **阅读状态:** 已载入记忆。")
    else:
        st.info("💡 **当前未加载任何文档。** 您可以直接向我咨询法规问题，或者在左侧边栏上传文档进行深度分析与修改。")

    # ==========================================
    # 6. Agent 聊天与核心调度逻辑
    # ==========================================
    if "messages" not in st.session_state:
        # 💥 核心修改区：赋予AI全能助手身份，同时保留专业能力
        system_prompt = """你是一个全能的AI智能助理，拥有广博的知识。但你的“核心专长”是资深医疗器械合规专家。

            你的工作模式必须遵循以下原则：
            1. 【常规闲聊与问答】：如果用户询问日常常识、天气、写诗、编程、翻译或其他非医疗领域的问题，请以热情、幽默、渊博的通用AI身份正常解答，绝对不要拒绝回答或局限于医疗领域。
            2. 【专业法规咨询】：如果用户询问医疗器械合规、MDR、质量体系等专业问题，请切换为严谨的专家身份进行解答。
            3. 【文档处理模式】：如果你需要基于用户上传的特定文档工作，你有两把武器可以调用：
               - search_document_content：阅读和搜索文档内容。
               - modify_word_document：把修改意见写回本地 Word。
            4.【公告机构NC整改】：当用户提到“公告机构”、“NC”、“整改方案”等关键词时，你必须切换到专家模式，结合专业知识和文档内容给出详细的分析和修改建议：
                请从以下审核报告提取不符合项，输出JSON格式：
                报告原文：...
                输出格式：[{"nc_id":"...", "description":"...", "clause":"...", "action":"..."}]
                每次NC整改生成可追溯的Excel矩阵：当用户要求生成整改矩阵时，你必须输出一个包含NC ID、描述、相关法规条款、整改措施和责任人的Excel表格，并提供下载链接。
            """
        st.session_state.messages = [
            SystemMessage(content=system_prompt),
            AIMessage(
                content="你好！我是一个全能AI助理，擅长医疗器械法规，也能陪你天南海北地闲聊。请问今天想聊点什么，或者有什么文档需要我帮忙处理？")
        ]

    if not api_key:
        st.stop()

    llm = ChatOpenAI(api_key=api_key, base_url="https://api.deepseek.com", model="deepseek-chat", temperature=0.1)
    llm_with_tools = llm.bind_tools(list(AVAILABLE_TOOLS.values()))

    for msg in st.session_state.messages:
        if isinstance(msg, SystemMessage) or isinstance(msg, ToolMessage):
            continue
        if isinstance(msg, AIMessage) and not msg.content:
            continue

        role = "user" if isinstance(msg, HumanMessage) else "assistant"
        with st.chat_message(role):
            st.write(msg.content)

    # 处理用户输入 (移除了没上传文档就不让聊天的限制)
    if user_input := st.chat_input("输入你的问题，或针对已上传文档的修改指令..."):
        with st.chat_message("user"):
            st.write(user_input)
        st.session_state.messages.append(HumanMessage(content=user_input))

        with st.chat_message("assistant"):
            with st.spinner("AI 正在思考..."):

                response = llm_with_tools.invoke(st.session_state.messages)
                st.session_state.messages.append(response)

                while response.tool_calls:
                    for tool_call in response.tool_calls:
                        tool_name = tool_call["name"]
                        tool_args = tool_call["args"]
                        tool_func = AVAILABLE_TOOLS.get(tool_name)

                        try:
                            if tool_func:
                                result_msg = tool_func.invoke(tool_args)
                            else:
                                result_msg = f"系统错误：找不到名为 {tool_name} 的工具。"
                        except Exception as e:
                            result_msg = f"工具执行异常失败: {str(e)}"

                        tool_message = ToolMessage(content=str(result_msg), tool_call_id=tool_call["id"])
                        st.session_state.messages.append(tool_message)

                    response = llm_with_tools.invoke(st.session_state.messages)
                    st.session_state.messages.append(response)

            if response.content:
                st.write(response.content)
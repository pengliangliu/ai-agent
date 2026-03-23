import os
import io
import datetime
import json
import base64
import pandas as pd
import fitz
import streamlit as st
from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage, ToolMessage
from langchain_core.tools import tool
from langchain_core.documents import Document as LangchainDocument
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from ddgs import DDGS

# ==========================================
# 0. 核心配置区
# ==========================================
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

# ==========================================
# 1. 页面配置
# ==========================================
st.set_page_config(page_title="全能法规智能助理", page_icon="🤖", layout="wide")


# ==========================================
# 2. 轻量级身份验证系统
# ==========================================
def check_password():
    USER_CREDENTIALS = {"admin": "123456", "boss": "888888"}

    def password_entered():
        username = st.session_state["login_username"]
        password = st.session_state["login_password"]
        api_key_input = st.session_state["login_api_key"].strip()

        if username in USER_CREDENTIALS and USER_CREDENTIALS[username] == password:
            if not api_key_input:
                st.session_state["password_correct"] = False
                st.session_state["login_error"] = "🚫 请输入有效的 DeepSeek API Key！"
            else:
                st.session_state["password_correct"] = True
                st.session_state["deepseek_api_key"] = api_key_input
                del st.session_state["login_password"]
                if "login_error" in st.session_state:
                    del st.session_state["login_error"]
        else:
            st.session_state["password_correct"] = False
            st.session_state["login_error"] = "🚫 用户名或密码错误，请重试！"

    if "password_correct" not in st.session_state or not st.session_state["password_correct"]:
        st.title("🔒 欢迎访问全能智能 Agent")
        st.info("请输入您的账号、密码以及 DeepSeek API Key 以继续。")
        st.text_input("用户名", key="login_username")
        st.text_input("密码", type="password", key="login_password")
        st.text_input("DeepSeek API Key (sk-...)", type="password", key="login_api_key")
        st.button("登录", on_click=password_entered, type="primary")
        if "login_error" in st.session_state:
            st.error(st.session_state["login_error"])
        return False
    return True


# ==========================================
# 🌟 核心拦截器
# ==========================================
if check_password():

    # ==========================================
    # 3. 性能优化与工具定义
    # ==========================================
    @st.cache_resource
    def load_embedding_model():
        return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")


    @tool
    def process_document_revision(action: str, original_text: str, revised_text: str, comment: str) -> str:
        """【智能文档修改与批注工具】当用户同意修改方案时调用！
        参数说明：
        - action (str): 填 "replace" 或 "append"。
        - original_text (str): 需要修改的原文片段（处理PDF时，请尽量提取短小精悍的关键词，避免使用包含换行的超长段落，否则会匹配失败）。
        - revised_text (str): 修改后的新文本。
        - comment (str): 给用户的说明。
        """
        if "current_file_path" not in st.session_state or not st.session_state.current_file_path:
            return "操作失败：当前没有加载任何文档。"

        source_path = st.session_state.current_file_path
        ext = os.path.splitext(source_path)[1].lower()

        try:
            # 🌟 模式一：处理 Word 文档 (.docx)
            if ext == '.docx':
                doc = DocxDocument(source_path)
                modified_count = 0

                if action == "append":
                    doc.add_paragraph("")
                    new_p = doc.add_paragraph()
                    run_text = new_p.add_run(revised_text)
                    run_text.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run_comment = new_p.add_run(f" [AI 追加章节: {comment}]")
                    run_comment.font.color.rgb = RGBColor(255, 0, 0)
                    modified_count += 1
                else:
                    for paragraph in doc.paragraphs:
                        if original_text and original_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(original_text, revised_text)
                            for run in paragraph.runs:
                                if revised_text in run.text:
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            run_comment = paragraph.add_run(f" [AI 批注: {comment}]")
                            run_comment.font.color.rgb = RGBColor(255, 0, 0)
                            modified_count += 1

                if modified_count > 0:
                    doc.save(source_path)
                    word_buffer = io.BytesIO()
                    doc.save(word_buffer)
                    b64_data = base64.b64encode(word_buffer.getvalue()).decode()

                    st.session_state.latest_modified_b64 = b64_data
                    st.session_state.latest_modified_ext = '.docx'
                    st.session_state.newly_modified_trigger = True
                    return f"成功：Word 文档已处理完毕。请回复：'文档已更新，请点击下方的蓝色按钮下载修订版。'"
                return f"失败：未在原文档中找到原文片段 '{original_text}'。"

            # 🌟 模式二：处理 PDF 文档 (.pdf)
            elif ext == '.pdf':
                if action == 'append':
                    return "操作提示：PDF 文件不支持直接追加章节，请告诉用户将内容写在备忘录里。"

                doc = fitz.open(source_path)
                modified_count = 0
                for page in doc:
                    text_instances = page.search_for(original_text)
                    if text_instances:
                        for inst in text_instances:
                            annot = page.add_highlight_annot(inst)
                            annot.set_info({"title": "AI 审核专家",
                                            "content": f"【建议修改为】:\n{revised_text}\n\n【说明】:\n{comment}"})
                            annot.update()
                            modified_count += 1
                    else:
                        # OCR 兜底模式
                        reader = load_ocr_model()
                        pix = page.get_pixmap(dpi=72)
                        ocr_results = reader.readtext(pix.tobytes("png"))

                        for bbox, text, prob in ocr_results:
                            if text.strip() and len(text.strip()) >= 2 and (
                                    text in original_text or original_text in text):
                                rect = fitz.Rect(bbox[0][0], bbox[0][1], bbox[2][0], bbox[2][1])
                                annot = page.add_highlight_annot(rect)
                                annot.set_info({"title": "AI 审核专家 (OCR视觉定位)",
                                                "content": f"【建议修改为】:\n{revised_text}\n\n【说明】:\n{comment}"})
                                annot.update()
                                modified_count += 1

                if modified_count > 0:
                    # 🚀 核心修复：将 PDF 压入内存字节，并立刻释放文件锁！
                    pdf_bytes = doc.write()
                    doc.close()

                    # 覆写本地临时文件，实现“叠甲修改”
                    with open(source_path, "wb") as f:
                        f.write(pdf_bytes)

                    b64_data = base64.b64encode(pdf_bytes).decode()

                    st.session_state.latest_modified_b64 = b64_data
                    st.session_state.latest_modified_ext = '.pdf'
                    st.session_state.newly_modified_trigger = True
                    return f"成功：已在 PDF 中高亮错误并批注。请回复：'PDF 已完成审核批注，请点击下方的红色按钮下载带有批注的 PDF。'"

                # 如果没找到，也要确保关掉文件，不然下次依然锁死
                doc.close()
                return f"失败：未在 PDF 中定位到原文 '{original_text}'。由于 PDF 存在隐形换行符，请你尝试提取短小精悍的连续词组作为 original_text，再次调用本工具重试！"

        except Exception as e:
            return f"文档处理时发生系统错误: {str(e)}"


    # 🌟 新增：从零开始生成全新 Word 文档的工具
    @tool
    def create_new_word_document(content: str, filename: str = "AI起草文档.docx") -> str:
        """【创建新Word文档工具】当用户要求你起草、撰写、生成一份全新的Word文档（且没有上传源文件）时调用此工具。
        参数说明：
        - content (str): 要写入新文档的完整文本内容（支持用 Markdown 的 #, ## 标识标题层级）。
        - filename (str): 生成的文件名，必须以 .docx 结尾。
        """
        try:
            doc = DocxDocument()

            # 简单的 Markdown 层级解析引擎，让生成的 Word 排版更好看
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line)

            if not filename.endswith('.docx'):
                filename += '.docx'

            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            b64_data = base64.b64encode(word_buffer.getvalue()).decode()

            # 存入纯内存状态
            st.session_state.latest_created_word_b64 = b64_data
            st.session_state.latest_created_word_filename = filename
            st.session_state.newly_created_word_trigger = True

            return f"成功：全新 Word 文档 {filename} 已生成。请回复：'文档已为您起草完毕，请点击下方的按钮直接下载。'"
        except Exception as e:
            return f"生成新文档时发生错误: {str(e)}"


    @tool
    def search_document_content(query: str) -> str:
        """【本地文档检索工具】当用户问起“文档里是怎么写的”时调用。"""
        if "vector_db" not in st.session_state or st.session_state.vector_db is None:
            return "操作失败：文档尚未被解析为向量数据库。"
        try:
            docs = st.session_state.vector_db.similarity_search(query, k=3)
            if not docs: return "检索完毕：没有找到相关内容。"
            context = "\n\n".join([f"原文段落 {i + 1}: {d.page_content}" for i, d in enumerate(docs)])
            return f"【本地检索结果】\n{context}"
        except Exception as e:
            return f"检索文档时发生错误: {str(e)}"


    @tool
    def search_latest_medical_regulations(query: str, time_limit: str = "m") -> str:
        """【互联网实时搜索工具】获取当前最新的法规、新闻。"""
        try:
            results = DDGS().text(query, max_results=3, timelimit=time_limit)
            if not results:
                return f"在 {time_limit} 范围内未能找到关于 '{query}' 的最新信息。"
            formatted_results = "\n\n".join(
                [f"标题: {res['title']}\n摘要: {str(res.get('body', ''))[:150]}...\n链接: {res['href']}" for res in
                 results])
            return f"【实时联网搜索结果】\n{formatted_results}"
        except Exception as e:
            return f"联网搜索发生异常: {str(e)}"


    @tool
    def generate_excel_matrix(json_data: str) -> str:
        """【Excel生成工具】生成NC整改矩阵。"""
        try:
            data = json.loads(json_data)
            df = pd.DataFrame(data)
            excel_buffer = io.BytesIO()
            df.to_excel(excel_buffer, index=False)
            b64_data = base64.b64encode(excel_buffer.getvalue()).decode()

            st.session_state.latest_excel_b64 = b64_data
            st.session_state.newly_generated_excel_trigger = True

            return "成功：Excel矩阵已在内存中生成。请回复：'表格已为您生成，请点击下方的绿色按钮直接下载。'"
        except Exception as e:
            return f"生成 Excel 时发生错误: {str(e)}。请检查传入的 JSON 格式。"


    @tool
    def get_file_download_link(file_type: str) -> str:
        """【获取文件下载链接】召唤历史生成的文件下载按钮。"""
        if file_type.lower() == 'doc':
            if st.session_state.get("latest_modified_b64"):
                st.session_state.newly_modified_trigger = True
                ext = st.session_state.latest_modified_ext
                doc_name = "Word" if ext == ".docx" else "PDF"
                return f"已触发 {doc_name} 下载按钮，请告诉用户：'文档链接已为您重新生成，请点击下方按钮获取。'"
            return "当前没有已修改的文档可供下载。"

        elif file_type.lower() == 'excel':
            if st.session_state.get("latest_excel_b64"):
                st.session_state.newly_generated_excel_trigger = True
                return "已触发 Excel 下载按钮，请告诉用户：'表格下载链接已为您重新生成，请点击下方绿色按钮获取。'"
            return "当前没有内存中的 Excel 表格可供下载。"

        return "未知的文件类型。请仅请求 doc 或 excel。"


    @tool
    def update_task_board(content: str) -> str:
        """【全局备忘录工具】当你分析长文档得出多个修改意见时，调用此工具将其记录在侧边栏，防止遗忘。"""
        st.session_state.task_board = content
        return "成功：已将内容安全写入全局备忘录。请停止执行并询问用户先处理备忘录里的哪一条。"


    AVAILABLE_TOOLS = {
        "process_document_revision": process_document_revision,
        "create_new_word_document": create_new_word_document,
        "search_document_content": search_document_content,
        "search_latest_medical_regulations": search_latest_medical_regulations,
        "generate_excel_matrix": generate_excel_matrix,
        "get_file_download_link": get_file_download_link,
        "update_task_board": update_task_board
    }


    def process_document_to_vector_db(file_path):
        docs = []
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            doc = fitz.open(file_path)
            for i, page in enumerate(doc):
                text = page.get_text("text")
                if len(text.strip()) > 5:
                    docs.append(LangchainDocument(page_content=text, metadata={"page": i + 1}))
        elif ext == '.docx':
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if len(p.text.strip()) > 5]
            docs = [LangchainDocument(page_content=text) for text in paragraphs]
        else:
            return None

        if not docs: return None
        embeddings = load_embedding_model()
        vector_db = FAISS.from_documents(docs, embeddings)
        return vector_db


    # ==========================================
    # 4. 状态初始化与主界面布局
    # ==========================================
    st.title("🤖 医疗器械法规 Agent")

    if "current_file_path" not in st.session_state: st.session_state.current_file_path = None
    if "current_file_name" not in st.session_state: st.session_state.current_file_name = None
    if "vector_db" not in st.session_state: st.session_state.vector_db = None
    if "task_board" not in st.session_state: st.session_state.task_board = ""

    if "latest_modified_b64" not in st.session_state: st.session_state.latest_modified_b64 = None
    if "latest_modified_ext" not in st.session_state: st.session_state.latest_modified_ext = None
    if "latest_excel_b64" not in st.session_state: st.session_state.latest_excel_b64 = None

    # 🌟 新增：凭空生成 Word 的状态
    if "latest_created_word_b64" not in st.session_state: st.session_state.latest_created_word_b64 = None
    if "latest_created_word_filename" not in st.session_state: st.session_state.latest_created_word_filename = None

    if "newly_generated_excel_trigger" not in st.session_state: st.session_state.newly_generated_excel_trigger = False
    if "newly_modified_trigger" not in st.session_state: st.session_state.newly_modified_trigger = False
    if "newly_created_word_trigger" not in st.session_state: st.session_state.newly_created_word_trigger = False

    # ==========================================
    # 5. 侧边栏
    # ==========================================
    with st.sidebar:
        st.header("⚙️ 系统状态")
        st.success("✅ 核心 AI 引擎已连接")
        st.markdown("---")
        st.header("📂 文档管理")
        uploaded_file = st.file_uploader("上传待处理文档 (.docx / .pdf)", type=["docx", "pdf"])

        if uploaded_file is not None:
            if st.session_state.current_file_name != uploaded_file.name:
                save_path = f"temp_{uploaded_file.name}"
                with open(save_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                st.session_state.current_file_path = save_path
                st.session_state.current_file_name = uploaded_file.name

                with st.spinner("正在逐字解析文档并建立大脑索引..."):
                    st.session_state.vector_db = process_document_to_vector_db(save_path)
                st.success(f"已加载并解析: {uploaded_file.name}")
        else:
            st.session_state.current_file_path = None
            st.session_state.current_file_name = None
            st.session_state.vector_db = None
            st.session_state.latest_modified_b64 = None
            st.session_state.latest_modified_ext = None
            st.session_state.latest_excel_b64 = None

        if st.session_state.latest_modified_b64:
            st.markdown("---")
            doc_bytes = base64.b64decode(st.session_state.latest_modified_b64)
            ext = st.session_state.latest_modified_ext
            filename = f"revised_{st.session_state.current_file_name}" if ext == ".docx" else f"annotated_{st.session_state.current_file_name}"
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == ".docx" else "application/pdf"
            st.download_button(f"📥 备用下载：处理后文档", data=doc_bytes, file_name=filename, mime=mime, type="primary")

        if st.session_state.task_board:
            st.markdown("---")
            st.header("📋 AI 待办备忘录")
            st.info(st.session_state.task_board)
            if st.button("🗑️ 清空备忘录"):
                st.session_state.task_board = ""
                st.rerun()

        st.markdown("---")
        if st.button("🚪 退出登录"):
            del st.session_state["password_correct"]
            if "deepseek_api_key" in st.session_state:
                del st.session_state["deepseek_api_key"]
            st.rerun()

    if st.session_state.current_file_name:
        st.info(f"📄 **当前文档:** `{st.session_state.current_file_name}` | 🧠 **阅读状态:** 已载入记忆。")
    else:
        st.info("💡 **当前未加载文档。** 您可以上传文档进行修改，也可以直接与我聊天")

    # ==========================================
    # 6. Agent 聊天与核心调度逻辑
    # ==========================================
    real_today = datetime.datetime.now().strftime("%Y年%m月%d日")
    doc_status = f"已在后台系统中加载并解析了文档【{st.session_state.current_file_name}】" if st.session_state.current_file_name else "当前未加载任何文档"
    current_board = st.session_state.task_board if st.session_state.task_board else "暂无内容"

    system_prompt = f"""你是一个全能的AI智能助理，核心专长是资深医疗器械合规专家。
        【重要时间认知】：今天是真实的 {real_today}。
        【当前文档状态】：{doc_status}
        【全局备忘录内容】：{current_board}

        🚫 【绝对禁令】（违背以下原则将被视为严重故障）：
        1. 你**绝对有能力**读取和修改用户上传的 Word 和 PDF 文件！当用户问“我上传的文件是干嘛的”或让你查阅文档时，你必须立刻调用 `search_document_content` 工具，**绝不允许**告诉用户“我无法接收文件附件”或“请复制粘贴文本”！
        2. 你的修改能力没有技术限制！**绝不允许**对用户说“由于系统技术限制需要手动修改”。

        你的工作模式：
        1. 【起草文档】：当用户没有上传文件要求起草文档时，调用 `create_new_word_document`。
        2. 【文档处理与分步策略】：发现文档有【多个错误】需要修改时：
           - 第一步：你必须先调用 `update_task_board` 工具，把所有需要整改的问题列入侧边栏备忘录。
           - 第二步：主动对用户说：“我已经把所有问题列在左侧备忘录了，为了保证准确性，咱们逐一修改，请问先从哪一条开始？”
           - 第三步：当用户指定修改某一条时，调用 `process_document_revision` 执行修改。改完后，询问用户是否继续修改备忘录里的下一条。
        3. 【生成表格】：调用 `generate_excel_matrix` 生成 Excel 文件。
        4. 【召唤链接】：调用 `get_file_download_link` 重新调出历史下载按钮。
        """

    if "messages" not in st.session_state:
        st.session_state.messages = [
            SystemMessage(content=system_prompt),
            AIMessage(
                content="你好！只要告诉我您的需求")
        ]
    else:
        if len(st.session_state.messages) > 0 and isinstance(st.session_state.messages[0], SystemMessage):
            st.session_state.messages[0] = SystemMessage(content=system_prompt)

    user_api_key = st.session_state.get("deepseek_api_key", "")
    llm = ChatOpenAI(api_key=user_api_key, base_url="https://api.deepseek.com", model="deepseek-chat", temperature=0.3,
                     streaming=True, max_retries=3, timeout=60.0)
    llm_with_tools = llm.bind_tools(list(AVAILABLE_TOOLS.values()))

    for msg in st.session_state.messages:
        if isinstance(msg, SystemMessage) or isinstance(msg, ToolMessage): continue
        if isinstance(msg, AIMessage) and not msg.content: continue
        role = "user" if isinstance(msg, HumanMessage) else "assistant"
        with st.chat_message(role):
            st.markdown(msg.content, unsafe_allow_html=True)

    if user_input := st.chat_input("输入指令..."):

        if len(st.session_state.messages) > 21:
            tail = st.session_state.messages[-20:]
            while tail and not isinstance(tail[0], HumanMessage):
                tail.pop(0)
            st.session_state.messages = [st.session_state.messages[0]] + tail

        if len(st.session_state.messages) > 0:
            last_msg = st.session_state.messages[-1]
            if isinstance(last_msg, AIMessage) and getattr(last_msg, "tool_calls", None):
                st.session_state.messages[-1] = AIMessage(content=last_msg.content)

        with st.chat_message("user"):
            st.markdown(user_input, unsafe_allow_html=True)
        st.session_state.messages.append(HumanMessage(content=user_input))

        with st.chat_message("assistant"):
            max_loops = 20
            current_loop = 0

            while current_loop < max_loops:
                current_loop += 1

                stop_btn_container = st.empty()
                message_placeholder = st.empty()
                ai_msg_chunk = None

                status_text = "AI 正在思考..." if current_loop == 1 else "AI 正在分析执行结果..."

                with stop_btn_container.container():
                    if st.button("⏹️ 停止生成", key=f"stop_{len(st.session_state.messages)}_{current_loop}"):
                        st.stop()

                with st.spinner(status_text):
                    current_messages = st.session_state.messages.copy()
                    st.session_state.messages.append(AIMessage(content=""))
                    current_msg_idx = len(st.session_state.messages) - 1

                    for chunk in llm_with_tools.stream(current_messages):
                        if ai_msg_chunk is None:
                            ai_msg_chunk = chunk
                        else:
                            ai_msg_chunk = ai_msg_chunk + chunk

                        st.session_state.messages[current_msg_idx] = ai_msg_chunk

                        if ai_msg_chunk.content:
                            message_placeholder.markdown(ai_msg_chunk.content + " ▌", unsafe_allow_html=True)

                if ai_msg_chunk and ai_msg_chunk.content:
                    message_placeholder.markdown(ai_msg_chunk.content, unsafe_allow_html=True)

                stop_btn_container.empty()

                if not ai_msg_chunk or not ai_msg_chunk.tool_calls:
                    break

                for tool_call in ai_msg_chunk.tool_calls:
                    tool_name = tool_call["name"]
                    tool_args = tool_call["args"]
                    tool_func = AVAILABLE_TOOLS.get(tool_name)

                    try:
                        if tool_func:
                            result_msg = tool_func.invoke(tool_args)
                        else:
                            result_msg = f"系统错误：找不到名为 {tool_name} 的工具。"
                    except Exception as e:
                        result_msg = f"工具执行异常: {str(e)}"

                    tool_message = ToolMessage(content=str(result_msg), tool_call_id=tool_call["id"])
                    st.session_state.messages.append(tool_message)

            if current_loop >= max_loops and ai_msg_chunk and ai_msg_chunk.tool_calls:
                st.warning("⚠️ 思考及搜索轮数达到上限，已强制中断推导。")

            # ==========================================
            # 🌟 拦截器：渲染纯内存下载按钮
            # ==========================================

            # 1. 拦截 Excel 下载
            if st.session_state.get("newly_generated_excel_trigger") and st.session_state.get("latest_excel_b64"):
                b64 = st.session_state.latest_excel_b64
                filename = "NC_Rectification_Matrix.xlsx"

                html_link = f'<div style="margin-top: 15px;"><a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{filename}" target="_blank" style="display: inline-block; padding: 10px 20px; background-color: #00A67E; color: white; text-align: center; text-decoration: none; border-radius: 6px; font-weight: bold;">📊 点击这里直接下载 Excel 矩阵表</a></div>'
                st.markdown(html_link, unsafe_allow_html=True)

                if len(st.session_state.messages) > 0 and isinstance(st.session_state.messages[-1], AIMessage):
                    st.session_state.messages[-1].content += "\n\n" + html_link

                st.session_state.newly_generated_excel_trigger = False

            # 2. 拦截修改过的源文档下载
            if st.session_state.get("newly_modified_trigger") and st.session_state.get("latest_modified_b64"):
                b64 = st.session_state.latest_modified_b64
                ext = st.session_state.latest_modified_ext

                if ext == '.docx':
                    filename = f"revised_{st.session_state.current_file_name}"
                    btn_color = "#0052CC"
                    btn_text = "📝 点击这里直接下载修订版 Word"
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    filename = f"annotated_{st.session_state.current_file_name}"
                    btn_color = "#E52B50"
                    btn_text = "📕 点击这里下载带有专家批注的 PDF"
                    mime_type = "application/pdf"

                html_link = f'<div style="margin-top: 15px;"><a href="data:{mime_type};base64,{b64}" download="{filename}" target="_blank" style="display: inline-block; padding: 10px 20px; background-color: {btn_color}; color: white; text-align: center; text-decoration: none; border-radius: 6px; font-weight: bold;">{btn_text}</a></div>'
                st.markdown(html_link, unsafe_allow_html=True)

                if len(st.session_state.messages) > 0 and isinstance(st.session_state.messages[-1], AIMessage):
                    st.session_state.messages[-1].content += "\n\n" + html_link

                st.session_state.newly_modified_trigger = False

            # 3. 🌟 新增：拦截全新起草的 Word 下载
            if st.session_state.get("newly_created_word_trigger") and st.session_state.get("latest_created_word_b64"):
                b64 = st.session_state.latest_created_word_b64
                filename = st.session_state.latest_created_word_filename

                # 用深蓝色与“修改源文件”的普通蓝色做轻微视觉区分
                html_link = f'<div style="margin-top: 15px;"><a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" target="_blank" style="display: inline-block; padding: 10px 20px; background-color: #003399; color: white; text-align: center; text-decoration: none; border-radius: 6px; font-weight: bold;">📄 点击这里下载新建的 Word 文档</a></div>'
                st.markdown(html_link, unsafe_allow_html=True)

                if len(st.session_state.messages) > 0 and isinstance(st.session_state.messages[-1], AIMessage):
                    st.session_state.messages[-1].content += "\n\n" + html_link

                st.session_state.newly_created_word_trigger = False